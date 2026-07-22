"""输出复核包生成器。

生成完整复核包（6 个文件），且不修改原件：
- 复核问题清单.xlsx：问题表，每行包含问题编号、报告层面、报表名称、问题类别、风险等级
- 逐项复核台账.xlsx：逐项台账，覆盖固定范围恰好一次（目标编号、检查范围、结论、状态）
- 未确认事项.xlsx：未确认事项清单（类型、描述、原因）
- 复核总结.docx：复核总结文档
- 证据索引.html：证据索引页，可从问题跳到原始节点
- 完成回执.json：保存输入、模板、模式、降级、隐藏测试、重试和输出摘要

设计要点：
- 生成前调用 project_service.verify_sources 检测源文件是否被篡改。
- 生成前要求所有目标已终态（_require_all_targets_terminal）。
- 不生成"修改后"文件，原件摘要不变。

ponytail: 当前实现为骨架版本，问题表/台账/未确认均只写表头，
证据索引为静态 HTML，回执 JSON 为占位结构。实际数据由后续任务
（问题汇聚、台账生成、回执汇总）接入后填充。
"""

import json
from pathlib import Path

import openpyxl


class OutputGenerator:
    """输出生成器。

    生成完整复核包（6 个文件），且不修改原件。

    Args:
        project_service: 项目服务，用于 verify_sources 检测源文件变化。
                         为 None 时跳过检测（用于无项目服务的场景，如测试）。
        output_dir: 输出目录，默认为 Path("output")。
    """

    def __init__(self, project_service=None, output_dir=None):
        self.project_service = project_service
        self.output_dir = output_dir or Path("output")

    def generate(self, project_id) -> tuple:
        """生成完整复核包。

        简报原样代码：
            self.project_service.verify_sources(project_id)
            self._require_all_targets_terminal(project_id)
            return self._write_package(project_id)

        适配：project_service 为 None 时跳过 verify_sources（用于测试）。

        Args:
            project_id: 项目 ID。

        Returns:
            6 个输出文件路径的元组。
        """
        if self.project_service:
            self.project_service.verify_sources(project_id)
        self._require_all_targets_terminal(project_id)
        return self._write_package(project_id)

    def _require_all_targets_terminal(self, project_id):
        """检查所有目标是否已终态。

        ponytail: 简化实现，总是通过。实际由后续任务接入 targets 表查询。
        """
        pass

    def _write_package(self, project_id) -> tuple:
        """写入 6 个输出文件，返回路径元组。"""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        files = (
            self._write_issues_xlsx(),
            self._write_ledger_xlsx(),
            self._write_unconfirmed_xlsx(),
            self._write_summary_docx(),
            self._write_evidence_html(),
            self._write_receipt_json(project_id),
        )
        return files

    def _write_issues_xlsx(self):
        """生成复核问题清单.xlsx。

        每行包含设计第 16.1 节字段：问题编号、报告层面、报表名称、问题类别、风险等级。
        ponytail: 当前仅写表头，实际数据由问题汇聚任务接入后填充。
        """
        path = self.output_dir / "复核问题清单.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "问题清单"
        ws.append(["问题编号", "报告层面", "报表名称", "问题类别", "风险等级"])
        wb.save(path)
        return path

    def _write_ledger_xlsx(self):
        """生成逐项复核台账.xlsx。

        逐项台账覆盖固定范围恰好一次：目标编号、检查范围、结论、状态。
        ponytail: 当前仅写表头，实际数据由台账生成任务接入后填充。
        """
        path = self.output_dir / "逐项复核台账.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "台账"
        ws.append(["目标编号", "检查范围", "结论", "状态"])
        wb.save(path)
        return path

    def _write_unconfirmed_xlsx(self):
        """生成未确认事项.xlsx。

        未确认事项清单：类型、描述、原因。
        ponytail: 当前仅写表头，实际数据由未确认事项汇聚任务接入后填充。
        """
        path = self.output_dir / "未确认事项.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "未确认"
        ws.append(["类型", "描述", "原因"])
        wb.save(path)
        return path

    def _write_summary_docx(self):
        """生成复核总结.docx。

        ponytail: 当前仅写标题与说明段，实际总结内容由总结生成任务接入后填充。
        """
        from docx import Document

        path = self.output_dir / "复核总结.docx"
        doc = Document()
        doc.add_heading("复核总结", 0)
        doc.add_paragraph("本报告由通用受控财务报表复核系统自动生成。")
        doc.save(path)
        return path

    def _write_evidence_html(self):
        """生成证据索引.html。

        证据索引可从问题跳到原始节点。
        ponytail: 当前为静态 HTML 骨架，实际锚点由证据索引任务接入后填充。
        """
        path = self.output_dir / "证据索引.html"
        path.write_text(
            "<html><head><title>证据索引</title></head>"
            "<body><h1>证据索引</h1></body></html>",
            encoding="utf-8",
        )
        return path

    def _write_receipt_json(self, project_id):
        """生成完成回执.json。

        回执保存输入、模板、模式、降级、隐藏测试、重试和输出摘要。
        ponytail: 当前为占位结构，实际数据由回执汇总任务接入后填充。
        """
        path = self.output_dir / "完成回执.json"
        receipt = {
            "project_id": project_id,
            "input_files": [],
            "template_version": "1.0",
            "quality_mode": "strict",
            "task_count": 0,
            "completion_status": "completed",
            "output_files": [],
        }
        path.write_text(
            json.dumps(receipt, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return path
